# modules/loader.py

import os
from langchain_community.document_loaders import PyPDFLoader
from langchain.schema import Document

DATA_FOLDER = "data"
SUPPORTED_FILE_TYPES = [".pdf", ".docx", ".png", ".jpg", ".jpeg"]


def is_supported_file(filename):
    return os.path.splitext(filename)[1].lower() in SUPPORTED_FILE_TYPES


def load_file(file_path):
    ext = os.path.splitext(file_path)[1].lower()
    documents = []

    try:
        if ext == ".pdf":
            loader = PyPDFLoader(file_path)
            documents = loader.load()

        elif ext == ".docx":
            try:
                from docx import Document as DocxDocument
            except ImportError as e:
                raise ImportError("python-docx is required to load DOCX files.") from e

            doc = DocxDocument(file_path)
            text = "\n".join([para.text for para in doc.paragraphs if para.text.strip()])
            if not text.strip():
                text = "[No text extracted from DOCX document.]"
            documents = [Document(page_content=text, metadata={"source": os.path.basename(file_path)})]

        elif ext in [".png", ".jpg", ".jpeg"]:
            try:
                from PIL import Image
            except ImportError as e:
                raise ImportError("Pillow is required to load image files.") from e

            try:
                import pytesseract
            except ImportError as e:
                raise ImportError("pytesseract is required to extract text from images.") from e

            with Image.open(file_path) as image:
                text = pytesseract.image_to_string(image)

            if not text.strip():
                text = "[No text extracted from image.]"
            documents = [Document(page_content=text, metadata={"source": os.path.basename(file_path)})]

        else:
            return []

        for doc in documents:
            if not getattr(doc, "metadata", None):
                doc.metadata = {}
            doc.metadata["source"] = os.path.basename(file_path)

        return documents

    except Exception as e:
        print(f"[ERROR] Failed to load {file_path}: {e}")
        return []


def load_all_documents():
    documents = []

    try:
        for file in os.listdir(DATA_FOLDER):
            if is_supported_file(file):
                file_path = os.path.join(DATA_FOLDER, file)
                docs = load_file(file_path)
                documents.extend(docs)

        print(f"[INFO] Loaded {len(documents)} pages from supported documents.")
        return documents

    except Exception as e:
        print(f"[ERROR] Failed to load documents from {DATA_FOLDER}: {e}")
        return None


load_all_pdfs = load_all_documents


